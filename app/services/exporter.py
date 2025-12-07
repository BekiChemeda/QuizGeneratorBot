import io
from typing import List, Dict
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document

class QuizExporter:
    @staticmethod
    def to_txt(title: str, questions: List[Dict]) -> io.BytesIO:
        output = io.BytesIO()
        text = f"Quiz: {title}\n\n"
        for i, q in enumerate(questions, 1):
            text += f"{i}. {q['question']}\n"
            for j, c in enumerate(q['choices']):
                idx_char = chr(65 + j)
                text += f"   {idx_char}. {c}\n"
            text += "\n"
        
        text += "\n--- ANSWERS ---\n"
        for i, q in enumerate(questions, 1):
             ans_char = chr(65 + q['answer_index'])
             text += f"{i}. {ans_char}\n"
             if q.get('explanation'):
                 text += f"   Explanation: {q['explanation']}\n"
        
        output.write(text.encode('utf-8'))
        output.seek(0)
        return output

    @staticmethod
    def to_pdf(title: str, questions: List[Dict]) -> io.BytesIO:
        output = io.BytesIO()
        c = canvas.Canvas(output, pagesize=letter)
        width, height = letter
        y = height - 40
        
        def check_page(curr_y):
            if curr_y < 50:
                c.showPage()
                return height - 40
            return curr_y

        c.setFont("Helvetica-Bold", 16)
        c.drawString(40, y, f"Quiz: {title}")
        y -= 30
        c.setFont("Helvetica", 12)

        for i, q in enumerate(questions, 1):
            y = check_page(y)
            # Wrap question text simple logic (or just truncate for now to avoid complexity without Paragraph)
            # Better: use simple textObject for multiline? Or just drawing string.
            # Using simple replacement for newlines if any
            q_text = f"{i}. {q['question']}"
            c.drawString(40, y, q_text)
            y -= 15
            
            for j, ch_text in enumerate(q['choices']):
                y = check_page(y)
                idx_char = chr(65 + j)
                c.drawString(60, y, f"{idx_char}. {ch_text}")
                y -= 15
            y -= 10
        
        # Answers Page
        c.showPage()
        y = height - 40
        c.setFont("Helvetica-Bold", 14)
        c.drawString(40, y, "Answer Key")
        y -= 20
        c.setFont("Helvetica", 10)
        
        for i, q in enumerate(questions, 1):
            y = check_page(y)
            ans_char = chr(65 + q['answer_index'])
            c.drawString(40, y, f"{i}. {ans_char}")
            if q.get('explanation'):
                 c.drawString(80, y, f"- {q['explanation'][:80]}...") # Truncate for simple PDF
            y -= 15
            
        c.save()
        output.seek(0)
        return output

    @staticmethod
    def to_docx(title: str, questions: List[Dict]) -> io.BytesIO:
        doc = Document()
        doc.add_heading(f"Quiz: {title}", 0)
        
        for i, q in enumerate(questions, 1):
            doc.add_paragraph(f"{i}. {q['question']}", style='List Number')
            for j, c in enumerate(q['choices']):
                idx_char = chr(65 + j)
                doc.add_paragraph(f"{idx_char}. {c}", style='List Bullet 2')
            doc.add_paragraph() # spacing

        doc.add_page_break()
        doc.add_heading("Answer Key", level=1)
        
        for i, q in enumerate(questions, 1):
            ans_char = chr(65 + q['answer_index'])
            p = doc.add_paragraph(f"{i}. {ans_char}")
            if q.get('explanation'):
                p.add_run(f"\nExplanation: {q['explanation']}").italic = True
        
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output
